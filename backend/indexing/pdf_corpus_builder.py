# -*- coding: utf-8 -*-
"""Build a structured RAG corpus from curated PDF / DOCX / DOC source documents."""
from __future__ import annotations

import json
import logging
import re
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Iterable, List, Optional

import pdfplumber

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None  # type: ignore[assignment,misc]

try:
    import win32com.client as win32com_client
    import pythoncom
except ImportError:
    win32com_client = None  # type: ignore[assignment]
    pythoncom = None  # type: ignore[assignment]

_SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc'}

from chunking.unified_chunker import UnifiedChunker
from indexing.professional_cleaning_pipeline import professional_cleaning_pipeline

logger = logging.getLogger(__name__)


@dataclass
class SectionBuffer:
    title: str
    level: int
    path: List[str]
    page_start: int
    page_end: int
    lines: List[str] = field(default_factory=list)

    def add_line(self, line: str, page_number: int) -> None:
        self.lines.append(line)
        self.page_end = page_number

    def render_text(self) -> str:
        return '\n'.join(self.lines).strip()


class PDFCorpusBuilder:
    """Extract structured chunks and metadata from curated curriculum PDFs."""

    HEADING_PATTERNS = [
        (re.compile(r'^第[一二三四五六七八九十百零\d]+[章节条款]'), 1),
        (re.compile(r'^[一二三四五六七八九十]+、'), 2),
        (re.compile(r'^\d+(?:\.\d+){0,3}[、.]?\s*'), 2),
        (re.compile(r'^[（(][一二三四五六七八九十\d]+[）)]'), 3),
    ]
    NOISE_PATTERNS = [
        re.compile(r'^\s*$'),
        re.compile(r'^\d+$'),
        re.compile(r'^[·•.．…\-_]{4,}$'),
    ]
    TABLE_HINTS = ('课程体系与毕业要求', '关联度矩阵', '支撑关系', '指标点', 'PBLString', '矩阵', '对应关系')
    EXCLUDED_FILE_KEYWORDS = ('学生手册',)
    PROGRAM_KEYWORDS = ('培养目标', '毕业要求', '课程体系', '学分要求', '实践教学', '培养方案', '学制', '授予学位')
    COURSE_SECTION_KEYWORDS = (
        '课程简介', '课程性质', '课程目标', '教学目标', '教学内容', '教学安排', '教学重点', '教学难点',
        '考核', '成绩评定', '实验', '实践', '课程设计', '学时分配', '先修课程', '教材', '参考资料',
        '教学大纲', '课程学时', '学分', '课程定位', '课程内容', '基本要求', '教学方法', '授课方式',
        '课程考核', '教学进度', '课程描述', '能力培养', '知识目标', '能力目标', '素质目标',
    )
    CODE_PATTERNS = [
        re.compile(r'\bpublic\s+class\b'),
        re.compile(r'\bSystem\.out\.println\b'),
        re.compile(r'\bimport\s+java\.'),
        re.compile(r'https?://'),
        re.compile(r'\b[A-Za-z_][A-Za-z0-9_]*\s*\('),
        re.compile(r'[{};<>]{3,}'),
    ]

    def __init__(
        self,
        source_root: Path,
        chunker: Optional[UnifiedChunker] = None,
        min_section_length: int = 120,
    ) -> None:
        self.source_root = Path(source_root)
        self.chunker = chunker or UnifiedChunker()
        self.min_section_length = min_section_length

    def discover_pdf_files(self) -> List[Path]:
        """Discover all supported document files (PDF, DOCX, DOC)."""
        files: List[Path] = []
        for ext in _SUPPORTED_EXTENSIONS:
            for path in sorted(self.source_root.rglob(f'*{ext}')):
                if self._include_file(path):
                    files.append(path)
        files.sort(key=lambda p: p.name)
        logger.info('Discovered %s curated document files under %s', len(files), self.source_root)
        return files

    def build_corpus(self, output_path: Path, pdf_files: Optional[List[Path]] = None) -> Dict:
        files = pdf_files or self.discover_pdf_files()
        documents: List[Dict] = []
        chunks: List[Dict] = []
        keyword_counter: Counter[str] = Counter()

        for pdf_path in files:
            document = self._extract_document(pdf_path)
            if not document:
                continue
            documents.append(document['metadata'])
            for chunk in document['chunks']:
                chunks.append(chunk)
                keyword_counter.update(chunk.get('keywords', []))

        corpus = {
            'metadata': {
                'version': '3.1',
                'source': 'enterprise_pdf_corpus_curated',
                'document_count': len(documents),
                'total_chunks': len(chunks),
                'chunking_method': 'pdf_structured_semantic_curated',
                'documents': documents,
                'keywords': [keyword for keyword, _ in keyword_counter.most_common(100)],
            },
            'chunks': chunks,
        }

        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(json.dumps(corpus, ensure_ascii=False, indent=2), encoding='utf-8')
        logger.info('Structured corpus written to %s (documents=%s chunks=%s)', output_path, len(documents), len(chunks))
        return corpus

    def _include_file(self, file_path: Path) -> bool:
        stem = file_path.stem
        if any(keyword in stem for keyword in self.EXCLUDED_FILE_KEYWORDS):
            return False
        if file_path.name == 'desktop.ini':
            return False
        if '人工智能专业培养方案+专业课教案材料' in str(file_path):
            return True
        return stem.startswith(('CSE', 'C2022', 'MTH')) or '培养方案' in stem or '教案' in stem or '大纲' in stem

    def _extract_document(self, file_path: Path) -> Optional[Dict]:
        document_id = self._slugify(file_path.stem)
        document_title = file_path.stem.strip()
        category = file_path.parent.name if file_path.parent != self.source_root else self.source_root.name
        document_kind = self._document_kind(document_title)
        base_metadata = {
            'document_id': document_id,
            'title': document_title,
            'source_path': str(file_path.relative_to(self.source_root.parent)).replace('\\', '/'),
            'category': category,
            'document_kind': document_kind,
        }

        suffix = file_path.suffix.lower()
        if suffix == '.pdf':
            lines_by_page = list(self._extract_lines(file_path))
        elif suffix == '.docx':
            lines_by_page = list(self._extract_lines_docx(file_path))
        elif suffix == '.doc':
            lines_by_page = list(self._extract_lines_doc(file_path))
        else:
            logger.warning('Unsupported file type: %s', file_path)
            return None
        if not lines_by_page:
            return None
        if document_kind == 'program_spec':
            lines_by_page = self._slice_program_lines(lines_by_page)
            if not lines_by_page:
                return None

        lines_by_page, page_report = professional_cleaning_pipeline.clean_pages(lines_by_page)
        if not lines_by_page:
            return None

        sections = self._collect_sections(lines_by_page, document_title)
        chunks: List[Dict] = []
        kept_sections = 0
        paragraph_count = 0
        for section_index, section in enumerate(sections):
            cleaned_lines = self._clean_section_lines(section.lines, document_kind=document_kind, title=section.title, path=section.path)
            paragraph_count += len(cleaned_lines)
            section_text = '\n'.join(cleaned_lines).strip()
            if not self._should_keep_section(section, section_text, document_kind):
                continue
            kept_sections += 1

            section_metadata = {
                **base_metadata,
                'document_type': document_kind,
                'section_title': section.title,
                'section_path': section.path,
                'section_path_text': ' > '.join(section.path),
                'section_level': section.level,
                'page_start': section.page_start,
                'page_end': section.page_end,
                'chunk_type': self._infer_chunk_type(section.title, section.path),
            }
            section_chunks = self.chunker.chunk_text(section_text, metadata=section_metadata)
            for chunk_index, chunk in enumerate(section_chunks):
                chunk_metadata = dict(chunk.get('metadata') or {})
                chunk_text = chunk['text']
                if self._is_low_signal_chunk(chunk_text, document_kind):
                    continue
                chunk_id = f'{document_id}::{section_index:03d}::{chunk_index:03d}'
                keywords = self._extract_keywords(chunk_text)
                chunks.append(
                    {
                        'id': chunk_id,
                        'text': chunk_text,
                        'char_count': len(chunk_text),
                        'section': ' > '.join(section.path),
                        'title': document_title,
                        'document_id': document_id,
                        'source_path': section_metadata['source_path'],
                        'page_start': section.page_start,
                        'page_end': section.page_end,
                        'keywords': keywords,
                        'chunk_type': section_metadata['chunk_type'],
                        'metadata': {**chunk_metadata, 'keywords': keywords},
                    }
                )

        chunks, quality_report = professional_cleaning_pipeline.assess_chunks(
            chunks,
            document_metadata={**base_metadata, 'document_type': document_kind},
            section_count=len(sections),
            paragraph_count=paragraph_count,
            page_report=page_report,
        )

        return {
            'metadata': {
                **base_metadata,
                'document_type': document_kind,
                'page_count': len(lines_by_page),
                'section_count': len(sections),
                'kept_section_count': kept_sections,
                'chunk_count': len(chunks),
                'quality_report': quality_report,
            },
            'chunks': chunks,
        }

    def _extract_lines(self, pdf_path: Path) -> Iterable[tuple[int, List[str]]]:
        with pdfplumber.open(pdf_path) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ''
                normalized_lines = []
                for line in text.splitlines():
                    cleaned = self._normalize_line(line)
                    if not cleaned or self._is_noise_line(cleaned) or self._is_toc_line(cleaned):
                        continue
                    normalized_lines.append(cleaned)
                if normalized_lines:
                    yield page_number, normalized_lines

    def _extract_lines_docx(self, docx_path: Path) -> Iterable[tuple[int, List[str]]]:
        """Extract lines from a .docx file using python-docx."""
        if DocxDocument is None:
            logger.warning('python-docx not installed, skipping %s', docx_path)
            return
        doc = DocxDocument(str(docx_path))
        page_number = 1
        page_lines: List[str] = []
        lines_per_page = 45  # approximate page break

        for para in doc.paragraphs:
            raw = para.text.strip()
            if not raw:
                continue
            cleaned = self._normalize_line(raw)
            if not cleaned or self._is_noise_line(cleaned) or self._is_toc_line(cleaned):
                continue
            page_lines.append(cleaned)
            if len(page_lines) >= lines_per_page:
                yield page_number, page_lines
                page_number += 1
                page_lines = []

        # Also extract from tables (common in teaching plans)
        for table in doc.tables:
            for row in table.rows:
                cell_texts = []
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in cell_texts:
                        cell_texts.append(text)
                if cell_texts:
                    merged = ' '.join(cell_texts)
                    cleaned = self._normalize_line(merged)
                    if cleaned and not self._is_noise_line(cleaned) and len(cleaned) > 5:
                        page_lines.append(cleaned)
                        if len(page_lines) >= lines_per_page:
                            yield page_number, page_lines
                            page_number += 1
                            page_lines = []

        if page_lines:
            yield page_number, page_lines

    def _extract_lines_doc(self, doc_path: Path) -> Iterable[tuple[int, List[str]]]:
        """Extract lines from a .doc file using win32com (Windows) or fallback."""
        if win32com_client is None:
            logger.warning('pywin32 not installed, skipping .doc file: %s', doc_path)
            return

        pythoncom.CoInitialize()
        word = None
        doc_obj = None
        text = ''
        try:
            word = win32com_client.DispatchEx('Word.Application')
            word.Visible = 0
            word.DisplayAlerts = 0
            abs_path = str(doc_path.resolve())
            doc_obj = word.Documents.Open(abs_path, ReadOnly=True, AddToRecentFiles=False)
            text = doc_obj.Content.Text or ''
        except Exception as exc:
            logger.error('Failed to open .doc via Word COM: %s — %s', doc_path, exc)
            text = ''
        finally:
            try:
                if doc_obj is not None:
                    doc_obj.Close(False)
            except Exception:
                pass
            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass
            pythoncom.CoUninitialize()

        if not text.strip():
            return

        page_number = 1
        page_lines: List[str] = []
        lines_per_page = 45

        for raw_line in text.splitlines():
            cleaned = self._normalize_line(raw_line.strip())
            if not cleaned or self._is_noise_line(cleaned) or self._is_toc_line(cleaned):
                continue
            page_lines.append(cleaned)
            if len(page_lines) >= lines_per_page:
                yield page_number, page_lines
                page_number += 1
                page_lines = []

        if page_lines:
            yield page_number, page_lines

    def _slice_program_lines(self, lines_by_page: List[tuple[int, List[str]]]) -> List[tuple[int, List[str]]]:
        major_heading_pattern = re.compile(r'.+专业(?:应用型)?人才培养方案$')
        start_index: Optional[int] = None
        end_index: Optional[int] = None
        flat_lines: List[tuple[int, str]] = []
        for page_number, lines in lines_by_page:
            for line in lines:
                flat_lines.append((page_number, line))

        for index, (_, line) in enumerate(flat_lines):
            if '人工智能专业' in line and major_heading_pattern.search(line):
                start_index = index
                break

        if start_index is None:
            return lines_by_page

        for index in range(start_index + 1, len(flat_lines)):
            _, line = flat_lines[index]
            if major_heading_pattern.search(line) and '人工智能专业' not in line:
                end_index = index
                break

        sliced = flat_lines[start_index:end_index]
        page_map: Dict[int, List[str]] = {}
        for page_number, line in sliced:
            page_map.setdefault(page_number, []).append(line)
        return [(page_number, page_map[page_number]) for page_number in sorted(page_map)]
    def _collect_sections(self, lines_by_page: Iterable[tuple[int, List[str]]], document_title: str) -> List[SectionBuffer]:
        sections: List[SectionBuffer] = []
        stack: List[str] = [document_title]
        current = SectionBuffer(title=document_title, level=0, path=[document_title], page_start=1, page_end=1)

        for page_number, lines in lines_by_page:
            for line in lines:
                heading = self._parse_heading(line)
                if heading:
                    if current.lines:
                        sections.append(current)
                    level, heading_title = heading
                    stack = self._merge_stack(stack, level, heading_title)
                    current = SectionBuffer(
                        title=heading_title,
                        level=level,
                        path=stack.copy(),
                        page_start=page_number,
                        page_end=page_number,
                    )
                    continue
                current.add_line(line, page_number)

        if current.lines:
            sections.append(current)
        return sections

    def _merge_stack(self, stack: List[str], level: int, title: str) -> List[str]:
        root = stack[:1] if stack else []
        body = stack[1:level]
        return root + body + [title]

    def _parse_heading(self, line: str) -> Optional[tuple[int, str]]:
        for pattern, level in self.HEADING_PATTERNS:
            if pattern.match(line):
                return level, line.strip(' ：:')
        return None

    def _normalize_line(self, line: str) -> str:
        normalized = re.sub(r'\s+', ' ', line.replace('\u3000', ' ')).strip()
        normalized = normalized.strip('|')
        normalized = normalized.replace('•', ' ')
        normalized = re.sub(r'^第\s+([一二三四五六七八九十百零\d]+)\s*([章节条款])', r'第\1\2', normalized)
        normalized = re.sub(r'^([一二三四五六七八九十]+)\s*[、.]\s*', r'\1、', normalized)
        return normalized

    def _clean_section_lines(self, lines: List[str], document_kind: str, title: str, path: List[str]) -> List[str]:
        cleaned: List[str] = []
        seen = set()
        for line in lines:
            line = self._normalize_line(line)
            if not line or line in seen:
                continue
            seen.add(line)
            if self._is_noise_line(line) or self._is_toc_line(line):
                continue
            if self._is_table_like_line(line):
                continue
            if document_kind == 'course_outline' and self._looks_like_code_or_runtime_line(line):
                continue
            if cleaned and self._should_merge_with_previous(cleaned[-1], line):
                cleaned[-1] = f"{cleaned[-1].rstrip()} {line.lstrip()}".strip()
                continue
            cleaned.append(line)
        return cleaned

    def _should_merge_with_previous(self, previous: str, current: str) -> bool:
        if not previous or not current:
            return False
        if self._parse_heading(current):
            return False
        if previous.endswith(('。', '！', '？', ';', '；', ':', '：')):
            return False
        if len(previous) >= 120:
            return False
        return True

    def _document_kind(self, title: str) -> str:
        if '培养方案' in title:
            return 'program_spec'
        return 'course_outline'

    def _should_keep_section(self, section: SectionBuffer, section_text: str, document_kind: str) -> bool:
        if len(section_text) < self.min_section_length:
            return False
        joined_path = ' '.join(section.path)
        if any(keyword in joined_path for keyword in self.TABLE_HINTS):
            return False
        if self._is_low_signal_chunk(section_text, document_kind):
            return False

        if document_kind == 'program_spec':
            return any(keyword in joined_path or keyword in section_text[:240] for keyword in self.PROGRAM_KEYWORDS)

        if any(keyword in joined_path or keyword in section_text[:240] for keyword in self.COURSE_SECTION_KEYWORDS):
            return True

        return section.level == 0 and any(keyword in section_text[:240] for keyword in self.COURSE_SECTION_KEYWORDS)

    def _is_low_signal_chunk(self, text: str, document_kind: str) -> bool:
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        ascii_letters = len(re.findall(r'[A-Za-z]', text))
        digits = len(re.findall(r'\d', text))
        total = max(len(text), 1)
        chinese_ratio = chinese_chars / total
        ascii_ratio = ascii_letters / total
        digit_ratio = digits / total
        if chinese_chars < 40:
            return True
        if ascii_letters > chinese_chars * 1.2:
            return True
        if chinese_ratio < 0.12 and ascii_ratio > 0.18:
            return True
        if digit_ratio > 0.35 and chinese_ratio < 0.2:
            return True
        if document_kind == 'course_outline' and any(pattern.search(text) for pattern in self.CODE_PATTERNS):
            return True
        return False

    def _is_noise_line(self, line: str) -> bool:
        if any(pattern.match(line) for pattern in self.NOISE_PATTERNS):
            return True
        if len(line) <= 2:
            return True
        if line.count('…') >= 3 or line.count('.') >= 8:
            return True
        if re.fullmatch(r'[\d\s\-_/]+', line):
            return True
        if '第 ' == line[:2]:
            return True
        return False

    def _is_table_like_line(self, line: str) -> bool:
        if any(hint in line for hint in self.TABLE_HINTS):
            return True
        number_tokens = re.findall(r'\b\d+(?:\.\d+)?\b', line)
        if len(number_tokens) >= 8 and len(re.findall(r'[\u4e00-\u9fff]', line)) < 20:
            return True
        if line.count('|') >= 3:
            return True
        if len(re.findall(r'[1-9]\s+[1-9]\s+[1-9]', line)) >= 1:
            return True
        return False

    def _looks_like_code_or_runtime_line(self, line: str) -> bool:
        if any(pattern.search(line) for pattern in self.CODE_PATTERNS):
            return True
        if line.startswith(('package ', 'class ', 'public ', 'private ', 'protected ')):
            return True
        if line.count('=') >= 2 and len(re.findall(r'[\u4e00-\u9fff]', line)) < 10:
            return True
        return False

    def _is_toc_line(self, line: str) -> bool:
        if re.search(r'[.．…]{4,}\s*\d+$', line):
            return True
        if '目录' in line and len(line) <= 10:
            return True
        return False

    def _infer_chunk_type(self, title: str, path: List[str]) -> str:
        joined = ' '.join(path)
        if any(keyword in joined for keyword in ['培养目标', '毕业要求', '课程体系', '学分要求', '实践教学']):
            return 'program_spec'
        if any(keyword in joined for keyword in ['课程目标', '教学内容', '考核', '实验', '实践', '课程性质', '课程简介']):
            return 'course_outline'
        return 'general'

    def _extract_keywords(self, text: str, top_n: int = 8) -> List[str]:
        import jieba

        stopwords = {'的', '了', '和', '及', '与', '在', '对', '并', '是', '为', '等', '要求', '课程', '专业'}
        tokens = [token.strip() for token in jieba.cut(text) if len(token.strip()) >= 2]
        counter = Counter(token for token in tokens if token not in stopwords)
        return [token for token, _ in counter.most_common(top_n)]

    def _slugify(self, value: str) -> str:
        slug = re.sub(r'[^0-9A-Za-z\u4e00-\u9fff]+', '-', value).strip('-')
        return slug.lower() or 'document'


